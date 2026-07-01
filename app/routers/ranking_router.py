import csv
import io
import logging
from fastapi import APIRouter, HTTPException, UploadFile, File, Request, Depends, status
from fastapi.responses import JSONResponse, StreamingResponse
from app.models.ranking_models import RankingRequest, RankingResponse, Candidate
from app.services.extraction_service import JDExtractionService
from app.services.ranking_service import CandidateRankingService
from app.services.candidate_loader import CandidateLoaderService
logger = logging.getLogger(__name__)
router = APIRouter(prefix='/api', tags=['Candidate Ranking'])

def get_extraction_service(request: Request) -> JDExtractionService:
    return request.app.state.extraction_service

def get_ranking_service(request: Request) -> CandidateRankingService:
    return request.app.state.ranking_service

def get_candidate_loader(request: Request) -> CandidateLoaderService:
    return request.app.state.candidate_loader

@router.post('/rank-candidates', response_model=RankingResponse, status_code=status.HTTP_200_OK, summary='Rank candidates against a job description', description='Full end-to-end pipeline. Accepts raw JD text + a list of candidate objects. Extracts hiring signals from the JD, then scores every candidate using a hybrid weighted scoring engine (no extra LLM calls = low latency). Returns a ranked shortlist with transparent score breakdowns.')
async def rank_candidates(request_body: RankingRequest, extraction_svc: JDExtractionService=Depends(get_extraction_service), ranking_svc: CandidateRankingService=Depends(get_ranking_service)):
    logger.info(f'POST /api/rank-candidates | candidates={len(request_body.candidates)} | jd_length={len(request_body.raw_jd_text)}')
    try:
        jd = await extraction_svc.extract(request_body.raw_jd_text)
    except ValueError as e:
        raise HTTPException(status_code=422, detail={'error': 'jd_extraction_failed', 'message': str(e)})
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail={'error': 'upstream_api_failure', 'message': str(e)})
    try:
        result = ranking_svc.rank(jd=jd, candidates=request_body.candidates, shortlist_size=request_body.shortlist_size or 10)
        return result
    except Exception as e:
        logger.exception(f'Ranking error: {e}')
        raise HTTPException(status_code=500, detail={'error': 'ranking_failed', 'message': str(e)})

@router.post('/rank-from-preloaded', response_model=RankingResponse, status_code=status.HTTP_200_OK, summary='Rank pre-loaded candidates against a JD', description='Uses the candidates pre-loaded from data/sample_candidates.json at startup. Only requires the raw JD text. Ideal for hackathon submission testing.')
async def rank_preloaded_candidates(raw_jd_text: str, shortlist_size: int=10, extraction_svc: JDExtractionService=Depends(get_extraction_service), ranking_svc: CandidateRankingService=Depends(get_ranking_service), candidate_loader: CandidateLoaderService=Depends(get_candidate_loader)):
    candidates = candidate_loader.get_candidates()
    if not candidates:
        raise HTTPException(status_code=404, detail={'error': 'no_candidates_loaded', 'message': 'No candidates found. Place sample_candidates.json in data/ directory.'})
    try:
        jd = await extraction_svc.extract(raw_jd_text)
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=502, detail=str(e))
    result = ranking_svc.rank(jd=jd, candidates=candidates, shortlist_size=shortlist_size)
    return result

@router.post('/rank-from-file', response_model=RankingResponse, status_code=status.HTTP_200_OK, summary='Upload JD file and rank pre-loaded candidates', description='Upload a .docx or .txt job description file. Extracts signals and ranks the pre-loaded sample_candidates.json dataset. This is the primary hackathon submission testing endpoint.')
async def rank_from_file(file: UploadFile=File(..., description='Upload the job_description.docx or .txt file'), shortlist_size: int=10, extraction_svc: JDExtractionService=Depends(get_extraction_service), ranking_svc: CandidateRankingService=Depends(get_ranking_service), candidate_loader: CandidateLoaderService=Depends(get_candidate_loader)):
    raw_text = ''
    try:
        content = await file.read()
        if file.filename.endswith('.txt'):
            raw_text = content.decode('utf-8', errors='ignore')
        elif file.filename.endswith('.docx'):
            import docx, io as _io
            doc = docx.Document(_io.BytesIO(content))
            raw_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise HTTPException(status_code=400, detail={'error': 'unsupported_file_type', 'message': 'Use .docx or .txt'})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail={'error': 'file_read_error', 'message': str(e)})
    if len(raw_text.strip()) < 50:
        raise HTTPException(status_code=400, detail={'error': 'file_too_short'})
    candidates = candidate_loader.get_candidates()
    if not candidates:
        raise HTTPException(status_code=404, detail={'error': 'no_candidates_loaded', 'message': 'Place sample_candidates.json in data/'})
    jd = await extraction_svc.extract(raw_text)
    result = ranking_svc.rank(jd=jd, candidates=candidates, shortlist_size=shortlist_size)
    return result

@router.post('/generate-output-csv', status_code=status.HTTP_200_OK, summary='Generate hackathon submission CSV from JD text', description='Runs the full pipeline and returns a downloadable CSV file in the hackathon submission format. Pass the raw JD text in the request body as a plain string.')
async def generate_output_csv(raw_jd_text: str, extraction_svc: JDExtractionService=Depends(get_extraction_service), ranking_svc: CandidateRankingService=Depends(get_ranking_service), candidate_loader: CandidateLoaderService=Depends(get_candidate_loader)):
    candidates = candidate_loader.get_candidates()
    if not candidates:
        raise HTTPException(status_code=404, detail={'error': 'no_candidates_loaded'})
    jd = await extraction_svc.extract(raw_jd_text)
    result = ranking_svc.rank(jd=jd, candidates=candidates, shortlist_size=len(candidates))
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['rank', 'candidate_id', 'name', 'total_score', 'must_have_skills_score', 'experience_score', 'nice_to_have_score', 'behavioral_score', 'location_score', 'notice_period_score', 'disqualifier_penalty', 'matched_must_have_skills', 'missing_must_have_skills', 'disqualifiers_hit', 'recruiter_note'])
    for rc in result.shortlist:
        writer.writerow([rc.rank, rc.candidate_id, rc.name, rc.total_score, rc.score_breakdown.must_have_skills_score, rc.score_breakdown.experience_score, rc.score_breakdown.nice_to_have_score, rc.score_breakdown.behavioral_score, rc.score_breakdown.location_score, rc.score_breakdown.notice_period_score, rc.score_breakdown.disqualifier_penalty, '; '.join(rc.matched_must_have_skills), '; '.join(rc.missing_must_have_skills), '; '.join(rc.disqualifiers_hit), rc.recruiter_note])
    output.seek(0)
    return StreamingResponse(iter([output.getvalue()]), media_type='text/csv', headers={'Content-Disposition': 'attachment; filename=ranked_candidates_output.csv'})

@router.get('/candidates/count', status_code=status.HTTP_200_OK, summary='How many candidates are pre-loaded', tags=['Health'])
async def candidate_count(candidate_loader: CandidateLoaderService=Depends(get_candidate_loader)):
    return JSONResponse(content={'candidates_loaded': candidate_loader.count()})

@router.post('/extract-candidate', status_code=status.HTTP_200_OK, summary='Extract candidate data from a resume file', description='Upload a .docx or .txt resume file. Returns a Candidate JSON object compatible with /api/rank-candidates. Used by the frontend UI to parse individual resumes before batch ranking.')
async def extract_candidate_from_resume(file: UploadFile=File(..., description='Upload a .docx or .txt resume file')):
    import uuid
    import re
    try:
        content = await file.read()
        if file.filename.endswith('.txt'):
            raw_text = content.decode('utf-8', errors='ignore')
        elif file.filename.endswith('.docx'):
            import docx, io as _io
            doc = docx.Document(_io.BytesIO(content))
            raw_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise HTTPException(status_code=400, detail={'error': 'unsupported_file_type', 'message': 'Use .docx or .txt'})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail={'error': 'file_read_error', 'message': str(e)})
    if len(raw_text.strip()) < 30:
        raise HTTPException(status_code=400, detail={'error': 'resume_too_short'})
    stem = file.filename.rsplit('.', 1)[0]
    candidate_name = stem.replace('_', ' ').replace('-', ' ').title()
    common_skills = ['Python', 'Java', 'JavaScript', 'SQL', 'PyTorch', 'TensorFlow', 'NLP', 'Machine Learning', 'Deep Learning', 'FastAPI', 'Django', 'React', 'Node.js', 'Docker', 'Kubernetes', 'AWS', 'GCP', 'Azure', 'embeddings', 'vector databases', 'FAISS', 'Pinecone', 'LLM', 'RAG', 'transformers', 'scikit-learn', 'pandas', 'Spark', 'Kafka']
    raw_lower = raw_text.lower()
    extracted_skills = [{'name': skill, 'proficiency': 'intermediate'} for skill in common_skills if skill.lower() in raw_lower]
    candidate = {'candidate_id': f'resume-{uuid.uuid4().hex[:8]}', 'name': candidate_name, 'summary': raw_text[:500], 'profile': {'years_of_experience': None, 'location': None, 'willing_to_relocate': False, 'current_industry': None}, 'skills': extracted_skills, 'career_history': [], 'redrob_signals': None}
    logger.info(f"Extracted candidate from resume | name='{candidate_name}' | skills_found={len(extracted_skills)} | file={file.filename}")
    return JSONResponse(content=candidate)